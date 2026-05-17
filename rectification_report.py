"""
rectification_report.py  —  STANDALONE RECTIFICATION REPORT MODULE
====================================================================
A fully self-contained module that generates a Rectification Report
as a .docx file (in-memory bytes).

ARCHITECTURE:
    ┌─────────────────────────────────────────────────────────┐
    │  rectification_report.py  (this file — all-in-one)      │
    │                                                          │
    │  ┌──────────────────────────────────────────────────┐   │
    │  │  TABLE ENGINE  (embedded)                        │   │
    │  │  generate_table_image()  — weighbridge table JPG │   │
    │  │  fetch_row_by_token()    — pull row from df      │   │
    │  │  fetch_row_from_series() — pull from pd.Series   │   │
    │  │  apply_user_updates()    — overlay action values  │   │
    │  │  fetch_and_generate()    — combined helper        │   │
    │  └──────────────────────────────────────────────────┘   │
    │                                                          │
    │  ┌──────────────────────────────────────────────────┐   │
    │  │  DOCUMENT ENGINE  (embedded)                     │   │
    │  │  build_rr_docx()  — main builder, returns bytes  │   │
    │  │  _build_headers_footers()                        │   │
    │  │  _add_section_break()                            │   │
    │  │  _add_signature_block()                          │   │
    │  │  _add_date_para()                                │   │
    │  │  _add_narrative()                                │   │
    │  │  _patch_docx()   — post-build XML fix            │   │
    │  └──────────────────────────────────────────────────┘   │
    │                                                          │
    │  ┌──────────────────────────────────────────────────┐   │
    │  │  STREAMLIT TRIGGER  (optional)                   │   │
    │  │  render_rr_trigger()  — UI entry point           │   │
    │  └──────────────────────────────────────────────────┘   │
    └─────────────────────────────────────────────────────────┘

INTEGRATION (minimal example):
    from rectification_report import build_rr_docx, fetch_and_generate,
                                     apply_user_updates, fetch_row_by_token

    tbl1_jpg, _   = fetch_and_generate(df, token, source="online",
                                       is_table2=False, force_outnet_yellow=True)
    tbl2_jpg, _   = fetch_and_generate(df, token, source="online",
                                       is_table2=True, override_values=action_data)
    docx_bytes, filename = build_rr_docx(
        token        = token,
        rr_serial    = "0290",
        rr_line      = "Rectification Report No. RR/B-44/2026/0290",
        before_dict  = fetch_row_by_token(df, token, source="online"),
        after_dict   = apply_user_updates(before_dict, action_data),
        action_data  = action_data,
        arr_dt       = datetime(2026, 4, 29, 17, 14),
        rpt_dt       = datetime(2026, 4, 30, 17, 14),
        tbl1_jpg     = tbl1_jpg,
        tbl2_jpg     = tbl2_jpg,
    )
    with open(filename, "wb") as f:
        f.write(docx_bytes)

STREAMLIT INTEGRATION:
    from rectification_report import render_rr_trigger
    render_rr_trigger(token=token, source=source, df=df,
                      row=row, action_data=action_data)

LOGO FILES:
    Place logo_toa.jpg and logo_samsung.jpg in the same directory
    as this file, or set LOGO_TOA_PATH / LOGO_SAMSUNG_PATH at the
    top of the file to absolute paths.

DEPENDENCIES (auto-installed on first run):
    python-docx, lxml, Pillow, matplotlib, pandas, openpyxl
"""

from __future__ import annotations

# ── Auto-install missing packages ────────────────────────────────────────────
import subprocess, sys

def _ensure_packages():
    required = {
        "docx":       "python-docx",
        "lxml":       "lxml",
        "PIL":        "Pillow",
        "matplotlib": "matplotlib",
        "pandas":     "pandas",
        "openpyxl":   "openpyxl",
    }
    for import_name, pip_name in required.items():
        try:
            __import__(import_name)
        except ImportError:
            subprocess.run(
                [sys.executable, "-m", "pip", "install", pip_name, "--quiet"],
                check=False,
            )

_ensure_packages()
# ─────────────────────────────────────────────────────────────────────────────

import io, math, os, re, tempfile, warnings
from datetime import datetime, timedelta
from pathlib import Path

import pandas as pd
import lxml.etree as ET

warnings.filterwarnings("ignore")

# ── Logo file paths — adjust if logos live elsewhere ─────────────────────────
_THIS_DIR       = Path(__file__).resolve().parent
LOGO_TOA_PATH    = _THIS_DIR / "logo_toa.jpg"
LOGO_SAMSUNG_PATH = _THIS_DIR / "logo_samsung.jpg"

# ── Month-to-B-code mapping (B-41 = Jan … B-52 = Dec) ───────────────────────
_MONTH_B = {1:41,2:42,3:43,4:44,5:45,6:46,7:47,8:48,9:49,10:50,11:51,12:52}

# ── Internal image counter (keeps Word relationship IDs unique) ───────────────
_IMG_CTR = [0]


# ══════════════════════════════════════════════════════════════════════════════
#  TABLE ENGINE  —  weighbridge data → JPEG table image
# ══════════════════════════════════════════════════════════════════════════════

# ── Column definitions ────────────────────────────────────────────────────────
RR_COLUMNS = [
    "TOKEN", "SITE CODE", "DATETIME ARRIVAL", "VEHICLE NO",
    "MATERIAL", "ACCEPTED", "IN WEIGHT", "OUT WEIGHT", "NET WEIGHT",
]

# Column widths in inches — proportional to expected content length
_COL_W = {
    "TOKEN":            2.8,
    "SITE CODE":        1.4,
    "DATETIME ARRIVAL": 2.0,
    "VEHICLE NO":       1.3,
    "MATERIAL":         1.4,
    "ACCEPTED":         1.1,
    "IN WEIGHT":        1.1,
    "OUT WEIGHT":       1.1,
    "NET WEIGHT":       1.1,
}

# Two-line header labels shown inside each column header cell
_HDR = {
    "TOKEN":            "TOKEN",
    "SITE CODE":        "SITE\nCODE",
    "DATETIME ARRIVAL": "DATE TIME\nARRIVAL",
    "VEHICLE NO":       "VEHICLE\nNO",
    "MATERIAL":         "MATERIAL",
    "ACCEPTED":         "ACCEPTED",
    "IN WEIGHT":        "IN\nWEIGHT",
    "OUT WEIGHT":       "OUT\nWEIGHT",
    "NET WEIGHT":       "NET\nWEIGHT",
}

# Column name maps for Online (default) data source
ONLINE_COL_MAP: dict[str, str] = {
    "TOKEN":            "Token",
    "SITE CODE":        "Source Site",
    "DATETIME ARRIVAL": "WB In Time",
    "VEHICLE NO":       "Vehicle Number",
    "MATERIAL":         "Material",
    "ACCEPTED":         "Accepted",
    "IN WEIGHT":        "In Weight",
    "OUT WEIGHT":       "Out Weight",
    "NET WEIGHT":       "Net Weight",
}

# Column name maps for Weighbridge data source
WB_COL_MAP: dict[str, str] = {
    "TOKEN":            "E-Token",
    "SITE CODE":        "Source Site",
    "DATETIME ARRIVAL": "Date Time Arrival",
    "VEHICLE NO":       "Vehicle No",
    "MATERIAL":         "Material",
    "ACCEPTED":         "Accepted",
    "IN WEIGHT":        "In Weight",
    "OUT WEIGHT":       "Out Weight",
    "NET WEIGHT":       "Net Weight",
}

# Colour palette
_HDR_BLUE = "#1F5C99"
_WHITE    = "#FFFFFF"
_BLACK    = "#000000"
_YELLOW   = "#FFFF00"
_RED      = "#CC0000"
_FONT     = "DejaVu Sans"


def _cell_colors(col: str, is_table2: bool, accepted_status: str,
                 force_outnet_yellow: bool = False) -> tuple[str, str]:
    """
    Return (background_hex, foreground_hex) for a data cell.

    Table 1 (is_table2=False, force_outnet_yellow=True):
        OUT WEIGHT, NET WEIGHT → Yellow background, Red text
        All other cells        → White background, Black text

    Table 2 (is_table2=True):
        ACCEPTED = "NO"        → Yellow background, Red text
        OUT WEIGHT             → Yellow background, Red text  (always)
        NET WEIGHT             → Yellow background, Red text  (always)
        All other cells        → White background, Black text
    """
    if force_outnet_yellow and col in ("OUT WEIGHT", "NET WEIGHT"):
        return _YELLOW, _RED
    if is_table2:
        if col == "ACCEPTED" and accepted_status == "NO":
            return _YELLOW, _RED
        if col in ("OUT WEIGHT", "NET WEIGHT"):
            return _YELLOW, _RED
    return _WHITE, _BLACK


def generate_table_image(
    row_dict: dict,
    is_table2: bool = False,
    dpi: int = 300,
    force_outnet_yellow: bool = False,
) -> bytes:
    """
    Render one weighbridge data row as a high-quality JPEG table image.

    Parameters
    ----------
    row_dict            : dict mapping RR_COLUMNS keys → display values
    is_table2           : True for the "after action" (corrected) table
    dpi                 : image resolution (300 recommended for print)
    force_outnet_yellow : if True, OUT/NET cells are yellow even in Table 1

    Returns
    -------
    bytes  — JPEG image data
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    from matplotlib import rcParams
    from PIL import Image

    rcParams["font.family"]    = _FONT
    rcParams["font.size"]      = 11
    rcParams["axes.linewidth"] = 0.8

    accepted_status = str(row_dict.get("ACCEPTED", "")).upper()

    col_widths = [_COL_W[c] for c in RR_COLUMNS]
    total_w    = sum(col_widths)
    hdr_h_in   = 0.72      # header row height in inches (room for 2-line labels)
    dat_h_in   = 0.58      # data row height in inches

    fig, ax = plt.subplots(figsize=(total_w, hdr_h_in + dat_h_in), dpi=dpi)
    ax.set_xlim(0, total_w)
    ax.set_ylim(0, hdr_h_in + dat_h_in)
    ax.axis("off")
    fig.patch.set_facecolor(_WHITE)

    dat_y0 = 0.0
    dat_y1 = dat_h_in
    hdr_y0 = dat_h_in
    hdr_y1 = hdr_h_in + dat_h_in

    x_pos = 0.0
    for ci, col in enumerate(RR_COLUMNS):
        cw = col_widths[ci]

        # Header cell
        ax.add_patch(plt.Rectangle(
            (x_pos, hdr_y0), cw, hdr_y1 - hdr_y0,
            facecolor=_HDR_BLUE, edgecolor="#222222", linewidth=0.7,
            transform=ax.transData, clip_on=False,
        ))
        ax.text(
            x_pos + cw / 2, (hdr_y0 + hdr_y1) / 2,
            _HDR[col],
            ha="center", va="center",
            fontsize=11, fontweight="bold",
            color=_WHITE, fontfamily=_FONT,
            linespacing=1.35, multialignment="center",
            transform=ax.transData, clip_on=False,
        )

        # Data cell
        bg, fg = _cell_colors(col, is_table2, accepted_status,
                              force_outnet_yellow=force_outnet_yellow)
        ax.add_patch(plt.Rectangle(
            (x_pos, dat_y0), cw, dat_y1 - dat_y0,
            facecolor=bg, edgecolor="#333333", linewidth=0.7,
            transform=ax.transData, clip_on=False,
        ))
        val = str(row_dict.get(col, "") or "")
        ax.text(
            x_pos + cw / 2, (dat_y0 + dat_y1) / 2,
            val,
            ha="center", va="center",
            fontsize=11, fontweight="bold",
            color=fg, fontfamily=_FONT,
            linespacing=1.35,
            transform=ax.transData, clip_on=False,
        )
        x_pos += cw

    # Outer border + header/data divider line
    ax.add_patch(plt.Rectangle(
        (0, 0), total_w, hdr_h_in + dat_h_in,
        fill=False, edgecolor="#111111", linewidth=1.5,
        transform=ax.transData, clip_on=False,
    ))
    ax.plot([0, total_w], [dat_y1, dat_y1],
            color="#333333", linewidth=0.9, transform=ax.transData)

    plt.subplots_adjust(left=0, right=1, top=1, bottom=0)

    # Save as lossless PNG first, then convert to JPEG (quality=97)
    buf_png = io.BytesIO()
    plt.savefig(buf_png, format="png", dpi=dpi,
                bbox_inches="tight", facecolor=_WHITE, edgecolor="none")
    plt.close(fig)
    buf_png.seek(0)

    from PIL import Image
    img = Image.open(buf_png).convert("RGB")
    buf_jpg = io.BytesIO()
    img.save(buf_jpg, format="JPEG", quality=97, dpi=(dpi, dpi), subsampling=0)
    return buf_jpg.getvalue()


# ── Row-fetch helpers ─────────────────────────────────────────────────────────

def _val_str(v) -> str:
    """Safely convert a raw DataFrame cell value to a display string."""
    if v is None or (isinstance(v, float) and math.isnan(v)):
        return ""
    if isinstance(v, (datetime, pd.Timestamp)):
        return pd.Timestamp(v).strftime("%d/%m/%Y %H:%M")
    return str(v).strip()


def fetch_row_by_token(
    df: pd.DataFrame,
    token: str,
    source: str = "online",
) -> dict:
    """
    Find the row matching *token* in *df* and return it mapped to RR_COLUMNS.

    Parameters
    ----------
    df     : the main weighbridge / online DataFrame
    token  : the E-Token string to search for
    source : "online" (default) or "wb"

    Returns
    -------
    dict  {RR_COLUMN_NAME: display_string, …}
    """
    col_map   = ONLINE_COL_MAP if source == "online" else WB_COL_MAP
    token_col = col_map["TOKEN"]

    if df is None or df.empty:
        raise ValueError(f"DataFrame is empty — cannot find token [{token}]")
    if token_col not in df.columns:
        raise ValueError(
            f"Token column '{token_col}' not found. "
            f"Available columns: {list(df.columns)}"
        )

    mask = df[token_col].astype(str).str.strip() == str(token).strip()
    if not mask.any():
        raise ValueError(f"E-Token [{token}] not found in column '{token_col}'")

    # Use first matching row for primary RR data (all rows available via fetch_all_rows_by_token)
    row    = df[mask].iloc[0]
    mapped = {col: _val_str(row.get(col_map[col], "")) for col in RR_COLUMNS}
    if mapped.get("ACCEPTED"):
        mapped["ACCEPTED"] = mapped["ACCEPTED"].upper()
    return mapped


def fetch_all_rows_by_token(
    df: pd.DataFrame,
    token: str,
    source: str = "online",
) -> list[dict]:
    """
    Return ALL rows matching *token* as a list of RR-column dicts.
    Used when a token has multiple heat entries.
    """
    col_map   = ONLINE_COL_MAP if source == "online" else WB_COL_MAP
    token_col = col_map["TOKEN"]

    if df is None or df.empty:
        return []
    if token_col not in df.columns:
        return []

    mask = df[token_col].astype(str).str.strip() == str(token).strip()
    rows = []
    for _, row in df[mask].iterrows():
        mapped = {col: _val_str(row.get(col_map[col], "")) for col in RR_COLUMNS}
        if mapped.get("ACCEPTED"):
            mapped["ACCEPTED"] = mapped["ACCEPTED"].upper()
        rows.append(mapped)
    return rows


def fetch_row_from_series(row: pd.Series, source: str = "online") -> dict:
    """
    Convert an existing pd.Series (one row already selected) to an RR dict.

    Parameters
    ----------
    row    : pd.Series — a single row from the weighbridge DataFrame
    source : "online" or "wb"

    Returns
    -------
    dict  {RR_COLUMN_NAME: display_string, …}
    """
    col_map = ONLINE_COL_MAP if source == "online" else WB_COL_MAP
    mapped  = {col: _val_str(row.get(col_map[col], "")) for col in RR_COLUMNS}
    if mapped.get("ACCEPTED"):
        mapped["ACCEPTED"] = mapped["ACCEPTED"].upper()
    return mapped


def apply_user_updates(base_row: dict, action_data: dict) -> dict:
    """
    Overlay operator-entered action values on top of the original row dict.

    action_data keys used: ACCEPTED, OUT WEIGHT, NET WEIGHT
    All other action_data keys are ignored (handled elsewhere).

    Returns a new dict — base_row is not mutated.
    """
    updated = base_row.copy()
    if "ACCEPTED" in action_data and action_data["ACCEPTED"] is not None:
        updated["ACCEPTED"] = str(action_data["ACCEPTED"]).upper()
    for key in ("OUT WEIGHT", "NET WEIGHT"):
        val = action_data.get(key, "")
        updated[key] = (
            str(val).strip()
            if val not in (None, "", "None", "nan")
            else ""
        )
    return updated


def fetch_and_generate(
    df: pd.DataFrame,
    token: str,
    source: str = "online",
    is_table2: bool = False,
    override_values: dict | None = None,
    dpi: int = 300,
    force_outnet_yellow: bool = False,
) -> tuple[bytes, dict]:
    """
    Combined convenience function: fetch row → optionally apply overrides
    → render table image.

    Returns
    -------
    (jpeg_bytes, row_dict)
    """
    row_dict = fetch_row_by_token(df, token, source)
    if override_values:
        row_dict = apply_user_updates(row_dict, override_values)
    jpg_bytes = generate_table_image(
        row_dict,
        is_table2=is_table2,
        dpi=dpi,
        force_outnet_yellow=force_outnet_yellow,
    )
    return jpg_bytes, row_dict


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT ENGINE  —  assembles the multi-page Rectification Report .docx
# ══════════════════════════════════════════════════════════════════════════════

def _sfx(d: int) -> str:
    """Ordinal suffix: 1→'st', 2→'nd', 3→'rd', else 'th'."""
    return {1: "st", 2: "nd", 3: "rd"}.get(
        d % 10 if d not in (11, 12, 13) else 0, "th"
    )


def _parse_dt(val) -> datetime:
    """Parse a date/datetime value from various formats; fallback = today."""
    if isinstance(val, (datetime, pd.Timestamp)):
        return pd.Timestamp(val).to_pydatetime()
    for fmt in (
        "%d/%m/%Y %H:%M", "%d/%m/%Y", "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d", "%d-%m-%Y %H:%M", "%d-%m-%Y",
    ):
        try:
            return datetime.strptime(str(val).strip(), fmt)
        except Exception:
            pass
    return datetime.today()


# ── Low-level Word XML helpers ────────────────────────────────────────────────

def _embed_image(doc, img_b: bytes, full_width: bool = False) -> None:
    """
    Inline-embed *img_b* (JPEG or PNG) into *doc* as a centred paragraph.

    full_width=False → fixed 6.3-inch width (table images on page 1)
    full_width=True  → full A4 usable width (~15.92 cm, page 4 screenshot)
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    from PIL import Image as PILImage

    _IMG_CTR[0] += 1
    n = _IMG_CTR[0]

    img = PILImage.open(io.BytesIO(img_b))
    iw, ih = img.size

    if full_width:
        PAGE_USABLE_EMU = int((21.0 - 2.54 - 2.54) / 2.54 * 914400)
        cx = PAGE_USABLE_EMU
    else:
        cx = int(6.3 * 914400)   # fixed original width — do NOT change

    cy = int(cx * ih / iw)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = __import__("docx").shared.Pt(0)
    p.paragraph_format.space_after  = __import__("docx").shared.Pt(4)

    is_png = img_b[:2] == b'\x89P'
    if is_png:
        uri  = PackURI(f"/word/media/rr_img_{n}.png")
        mime = "image/png"
    else:
        uri  = PackURI(f"/word/media/rr_tbl_{n}.jpg")
        mime = "image/jpeg"

    ip  = Part(uri, mime, img_b, p.part.package)
    rId = p.part.relate_to(
        ip,
        "http://schemas.openxmlformats.org/officeDocument/2006/"
        "relationships/image",
    )

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WP   = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    A    = "http://schemas.openxmlformats.org/drawingml/2006/main"
    PIC  = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    RL   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    xml = (
        f'<w:r xmlns:w="{W_NS}" xmlns:wp="{WP}" xmlns:a="{A}"'
        f' xmlns:pic="{PIC}" xmlns:r="{RL}">'
        f'<w:drawing><wp:inline><wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:docPr id="{n+600}" name="t{n}"/>'
        f'<a:graphic><a:graphicData uri="{PIC}">'
        f'<pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="t{n}"/>'
        f'<pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{rId}"/>'
        f'<a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr><a:xfrm><a:off x="0" y="0"/>'
        f'<a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>'
        f'</pic:pic></a:graphicData></a:graphic>'
        f'</wp:inline></w:drawing></w:r>'
    )
    p._p.append(ET.fromstring(xml))


def _add_logo_to_para(para, img_b: bytes, wcm: float, hcm: float) -> None:
    """Embed a logo image inside an existing paragraph (used in headers)."""
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    _IMG_CTR[0] += 1
    n   = _IMG_CTR[0]
    uri = PackURI(f"/word/media/hdr_logo_{n}.jpg")
    ip  = Part(uri, "image/jpeg", img_b, para.part.package)
    rId = para.part.relate_to(
        ip,
        "http://schemas.openxmlformats.org/officeDocument/2006/"
        "relationships/image",
    )

    cx = int(wcm * 914400 / 2.54)
    cy = int(hcm * 914400 / 2.54)

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WP   = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    A    = "http://schemas.openxmlformats.org/drawingml/2006/main"
    PIC  = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    RL   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    xml = (
        f'<w:r xmlns:w="{W_NS}" xmlns:wp="{WP}" xmlns:a="{A}"'
        f' xmlns:pic="{PIC}" xmlns:r="{RL}">'
        f'<w:drawing><wp:inline><wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:docPr id="{n+700}" name="hl_{n}"/>'
        f'<a:graphic><a:graphicData uri="{PIC}">'
        f'<pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="h{n}"/>'
        f'<pic:cNvPicPr/></pic:nvPicPr>'
        f'<pic:blipFill><a:blip r:embed="{rId}"/>'
        f'<a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
        f'<pic:spPr><a:xfrm><a:off x="0" y="0"/>'
        f'<a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'<a:ln><a:noFill/></a:ln>'
        f'</pic:spPr>'
        f'</pic:pic></a:graphicData></a:graphic>'
        f'</wp:inline></w:drawing></w:r>'
    )
    para._p.append(ET.fromstring(xml))


def _footer_page_numbers(ftr) -> None:
    """Add centred 'Page X of Y' auto-field to a footer object."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    fp = ftr.add_paragraph()
    fp.alignment = __import__("docx").enum.text.WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after  = Pt(0)

    def _fld(ins):
        for ft in ("begin", None, "separate", None, "end"):
            r = fp.add_run()
            if ft is None:
                if ins:
                    el = OxmlElement("w:instrText")
                    el.set(
                        "{http://www.w3.org/XML/1998/namespace}space",
                        "preserve",
                    )
                    el.text = f" {ins} "
                    r._r.append(el)
                    ins = None
                else:
                    r.text = "1"
                    r.font.size = Pt(9)
            else:
                fc = OxmlElement("w:fldChar")
                fc.set(
                    qn("w:fldCharType"),
                    {"begin": "begin", "separate": "separate", "end": "end"}[ft],
                )
                r._r.append(fc)

    fp.add_run("Page ").font.size = Pt(9)
    _fld("PAGE")
    fp.add_run(" of ").font.size = Pt(9)
    _fld("NUMPAGES")


# ── Document section builders ─────────────────────────────────────────────────

def _build_headers_footers(
    doc, rr_line: str, toa_bytes: bytes, sam_bytes: bytes
) -> None:
    """
    Configure page 1 header (title + logos + bottom border) and
    pages 2+ header (RR number underlined), plus footers (Page X of Y)
    on all pages.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    sec = doc.sections[0]
    sec.different_first_page_header_footer = True

    def _clear(obj):
        obj.is_linked_to_previous = False
        for p in list(obj.paragraphs):
            p._element.getparent().remove(p._element)

    # ── First page header: project title + logos + bottom separator ───────────
    fh = sec.first_page_header
    _clear(fh)

    tp = fh.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after  = Pt(4)
    tr = tp.add_run(
        "MANAGEMENT OF STAGING GROUND AND INFILLING WORKS (PHASE 3)"
    )
    tr.font.name = "Times New Roman"
    tr.font.size = Pt(12)
    tr.bold      = True
    tr.underline = True

    lp = fh.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after  = Pt(6)
    if toa_bytes:
        _add_logo_to_para(lp, toa_bytes, 4.2, round(4.2 * 71 / 326, 3))
    lp.add_run("     ")
    if sam_bytes:
        _add_logo_to_para(lp, sam_bytes, 5.0, round(5.0 * 81 / 441, 3))

    # Bottom border on the logo paragraph
    pPr  = lp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "000000")
    pBdr.append(bot)
    sp = pPr.find(qn("w:spacing"))
    if sp is not None:
        sp.addprevious(pBdr)
    else:
        pPr.append(pBdr)

    # ── Default header (pages 2+): RR number ─────────────────────────────────
    dh = sec.header
    _clear(dh)
    hp = dh.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after  = Pt(4)
    hr = hp.add_run(rr_line)
    hr.font.name = "Arial"
    hr.font.size = Pt(11)
    hr.underline = True

    # ── Footers: Page X of Y ──────────────────────────────────────────────────
    _clear(sec.first_page_footer)
    _footer_page_numbers(sec.first_page_footer)
    _clear(sec.footer)
    _footer_page_numbers(sec.footer)


def _add_date_para(doc, label: str, dt: datetime) -> None:
    """Add 'Date: 30th April 2026' with superscript ordinal suffix."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)

    def _r(text, sup=False):
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        if sup:
            va = OxmlElement("w:vertAlign")
            va.set(qn("w:val"), "superscript")
            run._r.get_or_add_rPr().append(va)

    _r(label)
    _r(str(dt.day))
    _r(_sfx(dt.day), sup=True)
    _r(f" {dt.strftime('%B %Y')}")


# ── Narrative reason codes ────────────────────────────────────────────────────
REASON_A = "Accepted Towing Vehicle"    # Option A – accepted, broke down AFTER unload
REASON_B = "Rejected Towing Vehicle"   # Option B – rejected, broke down BEFORE unload
REASON_C = "Late Time / Breakdown"     # Option C – late exit, breakdown after unload


def _add_narrative(
    doc,
    before_dict: dict,
    token:       str,
    arr_dt:      datetime,
    reason:      str = REASON_A,
) -> None:
    """
    Add the explanatory narrative paragraph to Page 1.

    reason=REASON_A  → Option A: Accepted Towing Vehicle
        Common header + accepted/towing text (broke down AFTER unload)
    reason=REASON_B  → Option B: Rejected Towing Vehicle
        Common header + rejected/towing text (broke down BEFORE unload)
    reason=REASON_C  → Option C: Late Time / Breakdown
        Common header + late-exit/mechanic text (broke down AFTER unload)

    All placeholder values are extracted dynamically from before_dict / arr_dt.
    Sentence structure and punctuation must NOT be changed.
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # ── Dynamic variable extraction ───────────────────────────────────────────
    truck       = str(before_dict.get("VEHICLE NO",  "[TRUCK_NO]")).strip()
    source_site = str(before_dict.get("SITE CODE",   "[SOURCE_SITE]")).strip()
    # Date formatted as: "29th April 2026 (Wednesday)"
    day_num   = arr_dt.day
    day_sfx   = _sfx(day_num)       # "st" / "nd" / "rd" / "th"
    month_yr  = arr_dt.strftime("%B %Y")
    weekday   = arr_dt.strftime("%A")

    # ── Paragraph setup ───────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)

    def _r(text, sup=False, bold=False):
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.bold      = bold
        if sup:
            va = OxmlElement("w:vertAlign")
            va.set(qn("w:val"), "superscript")
            run._r.get_or_add_rPr().append(va)

    # ── COMMON header: "Truck no. {TRUCK_NO} from source site {SOURCE_SITE}
    #    ({E-Token}), dated {DATE}" ───────────────────────────────────────────
    _r(f"Truck no. {truck} from source site {source_site} ({token}), dated ")
    _r(str(day_num))
    _r(day_sfx, sup=True)
    _r(f" {month_yr} ({weekday})")

    # ── Option-specific tail text (exact wording — do NOT modify) ─────────────
    reason_norm = str(reason).strip()

    if reason_norm == REASON_A:
        # Option A – Accepted Towing Vehicle
        _r(
            ", broke down after the material was unloaded at the SOFT CLAY"
            " platform. The truck was followed by a towing vehicle, so the"
            " ticket was not created because the transaction was not completed"
            " (please refer to the attached manual ticket). TSJV will make the"
            " necessary corrections on the affected source site and will inform"
            " Echol-Tech to update the system accordingly."
        )

    elif reason_norm == REASON_B:
        # Option B – Rejected Towing Vehicle
        _r(
            ", was rejected because it broke down before the material was"
            " unloaded at the SOFT CLAY platform. The truck was followed by a"
            " towing vehicle, so the ticket was not created as the transaction"
            " was not completed (please refer to the attached manual ticket)."
            " TSJV will make the necessary corrections on the affected source"
            " site and will inform Echol-Tech to update the system accordingly."
        )

    else:
        # Option C – Late Time / Breakdown (default fallback)
        # Sentence starts with "The truck, bearing no. …" — re-build paragraph
        # We already added the common prefix runs; clear and rebuild for Option C
        # which has a different sentence start structure.
        for child in list(p._p):
            p._p.remove(child)

        _r(f"The truck, bearing no. {truck} from source site {source_site}"
           f" ({token}), dated ")
        _r(str(day_num))
        _r(day_sfx, sup=True)
        _r(
            f" {month_yr} ({weekday}), broke down after the material was"
            " unloaded at the SOFT CLAY platform. The mechanic could not repair"
            " the vehicle in time, preventing it from exiting on schedule."
            " Consequently, the system could not close the transaction and"
            " flagged it for a manual check. TSJV will adjust necessary"
            " corrections on the affected source site and will inform"
            " Echol-Tech to update the system accordingly, as the transaction"
            " was not completed."
        )


def _add_section_break(doc, rr_line: str) -> None:
    """
    Insert a page break with a new section that has:
      - A4 page size / margins
      - Header showing the RR number (underlined)
      - Footer showing Page X of Y
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pb = doc.add_paragraph()
    pb.paragraph_format.space_before = Pt(0)
    pb.paragraph_format.space_after  = Pt(0)

    pPr  = pb._p.get_or_add_pPr()
    sP   = OxmlElement("w:sectPr")

    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), str(int(21.0 * 1440 / 2.54)))
    pgSz.set(qn("w:h"), str(int(29.7 * 1440 / 2.54)))
    sP.append(pgSz)

    pgMar = OxmlElement("w:pgMar")
    for k, v in [
        ("w:top", 2.5), ("w:right", 2.54), ("w:bottom", 2.0),
        ("w:left", 2.54), ("w:header", 1.27), ("w:footer", 1.27),
    ]:
        pgMar.set(qn(k), str(int(v * 1440 / 2.54)))
    pgMar.set(qn("w:gutter"), "0")
    sP.append(pgMar)
    pPr.append(sP)

    # Configure the new section's header and footer
    sec = doc.sections[-1]
    sec.different_first_page_header_footer = False

    hdr = sec.header
    hdr.is_linked_to_previous = False
    for p in list(hdr.paragraphs):
        p._element.getparent().remove(p._element)

    hp = hdr.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after  = Pt(4)
    hr = hp.add_run(rr_line)
    hr.font.name = "Arial"
    hr.font.size = Pt(11)
    hr.underline = True

    ftr = sec.footer
    ftr.is_linked_to_previous = False
    for p in list(ftr.paragraphs):
        p._element.getparent().remove(p._element)
    _footer_page_numbers(ftr)


def _add_signature_block(doc) -> None:
    """
    Two-column signature area at the bottom of Page 1:
        Name / Signature / Date    |    Name / Signature / Date
        TOA – Samsung C&T JV       |    Surbana Jurong Consultants
    Uses Word centre-tab stops to align the two columns.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    LEFT_CTR  = 2494    # twips from left margin
    RIGHT_CTR = 7007    # twips from left margin

    def _set_centre_tabs(para):
        pPr  = para._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs")
        for pos in (LEFT_CTR, RIGHT_CTR):
            t = OxmlElement("w:tab")
            t.set(qn("w:val"), "center")
            t.set(qn("w:pos"), str(pos))
            tabs.append(t)
        sp = pPr.find(qn("w:spacing"))
        if sp is not None:
            sp.addprevious(tabs)
        else:
            pPr.append(tabs)

    def _para(space_before=0, space_after=4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        _set_centre_tabs(p)
        return p

    def _tab(para):
        r = para.add_run()
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r._r.append(OxmlElement("w:tab"))

    def _run(para, text, underline=False):
        r = para.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.underline = underline

    # Underlined blank lines for signatures
    p0 = _para(space_before=0, space_after=1)
    _tab(p0); _run(p0, "\u00a0" * 34, underline=True)
    _tab(p0); _run(p0, "\u00a0" * 34, underline=True)

    # "Name / Signature / Date" labels
    p1 = _para(space_before=1, space_after=3)
    _tab(p1); _run(p1, "Name / Signature / Date")
    _tab(p1); _run(p1, "Name / Signature / Date")

    # Organisation names
    p2 = _para(space_before=2, space_after=4)
    _tab(p2); _run(p2, "TOA \u2013 Samsung C&T Joint Venture")
    _tab(p2); _run(p2, "Surbana Jurong Consultants Pte Ltd")


def _patch_docx(path: str) -> None:
    """
    Post-build XML patch applied to the saved .docx file:
      1. Ensure zoom is set to 100%.
      2. Propagate header/footer references from the last (body) section
         to inner sections so page numbers render correctly throughout.
      3. Add <w:titlePg/> to the first section so the different-first-page
         header/footer is honoured.
    """
    import zipfile, re as _re, io as _io

    with zipfile.ZipFile(path, "r") as z:
        names = z.namelist()
        sx = z.read("word/settings.xml").decode() if "word/settings.xml" in names else ""
        dx = z.read("word/document.xml").decode() if "word/document.xml" in names else ""

    # Fix zoom
    sx = _re.sub(
        r'<w:zoom(?![^>]*w:percent)([^/]*)/>', r'<w:zoom\1 w:percent="100"/>', sx
    )
    sx = sx.replace("<w:zoom/>", '<w:zoom w:percent="100"/>')

    all_m = list(_re.finditer(r'<w:sectPr[^>]*>.*?</w:sectPr>', dx, _re.DOTALL))
    if len(all_m) >= 2:
        body  = all_m[-1].group(0)
        first = all_m[0].group(0)
        hr    = _re.findall(r'<w:headerReference[^/]*/>', body)
        fr    = _re.findall(r'<w:footerReference[^/]*/>', body)
        def_h = next((r for r in hr if 'type="default"' in r), "")
        def_f = next((r for r in fr if 'type="default"' in r), "")
        refs  = "".join(hr) + "".join(fr)

        nf = (
            first.replace("<w:pgSz", refs + "<w:pgSz")
            if "<w:pgSz" in first
            else first.replace("</w:sectPr>", refs + "</w:sectPr>")
        )
        nf = (
            nf.replace("<w:cols",    "<w:titlePg/><w:cols")    if "<w:cols"    in nf else
            nf.replace("<w:docGrid", "<w:titlePg/><w:docGrid") if "<w:docGrid" in nf else
            nf.replace("</w:sectPr>", "<w:titlePg/></w:sectPr>")
        )

        nb = body
        for ref in hr + fr:
            nb = nb.replace(ref, "")
        nb = _re.sub(r'<w:titlePg[^/]*/>', '', nb)

        pat = dx.replace(first, nf, 1)
        bm  = list(_re.finditer(r'<w:sectPr[^>]*>.*?</w:sectPr>', pat, _re.DOTALL))
        if bm:
            pat = pat[: bm[-1].start()] + nb + pat[bm[-1].end():]

        for m in list(
            _re.finditer(r'<w:sectPr[^>]*>.*?</w:sectPr>', pat, _re.DOTALL)
        )[1:-1]:
            ob  = m.group(0)
            nb2 = _re.sub(r'<w:titlePg[^/]*/>', '', ob)
            if "headerReference" not in nb2 and def_h:
                nb2 = (
                    nb2.replace("<w:pgSz", def_h + def_f + "<w:pgSz")
                    if "<w:pgSz" in nb2
                    else nb2.replace("</w:sectPr>", def_h + def_f + "</w:sectPr>")
                )
            pat = pat.replace(ob, nb2, 1)
        dx = pat

    buf = _io.BytesIO()
    with zipfile.ZipFile(path, "r") as zin:
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == "word/settings.xml":
                    zout.writestr(item, sx.encode())
                elif item.filename == "word/document.xml":
                    zout.writestr(item, dx.encode())
                else:
                    zout.writestr(item, zin.read(item.filename))

    with open(path, "wb") as f:
        f.write(buf.getvalue())


# ══════════════════════════════════════════════════════════════════════════════
#  PUBLIC API: build_rr_docx
# ══════════════════════════════════════════════════════════════════════════════

def build_rr_docx(
    token:               str,
    rr_serial:           str,
    rr_line:             str,
    before_dict:         dict,
    after_dict:          dict,
    action_data:         dict,
    arr_dt:              datetime,
    rpt_dt:              datetime,
    tbl1_jpg:            bytes,
    tbl2_jpg:            bytes,
    excel_screenshot_jpg: "bytes | None" = None,
    weight_label:        str = "",
    reason:              str = REASON_A,
    filter_date_str:     str = "",    # YYYY-MM-DD from the UI date selector
) -> tuple[bytes, str]:
    """
    Build the complete Rectification Report as an in-memory .docx.

    Parameters
    ----------
    token               : E-Token identifier string
    rr_serial           : 4-digit serial string e.g. "0290"
    rr_line             : Full RR reference line
                          e.g. "Rectification Report No. RR/B-44/2026/0290"
    before_dict         : Row dict (original values) from fetch_row_*()
    after_dict          : Row dict with corrected values from apply_user_updates()
    action_data         : Raw action values dict; keys used:
                            ACCEPTED         — "YES" or "NO"
                            NET WEIGHT       — corrected net weight string
                            OUT WEIGHT       — corrected out weight string
                            UNLADEN WEIGHT   — unladen weight string (YES only)
    arr_dt              : Arrival datetime (from the weighbridge record)
    rpt_dt              : Report date (usually arr_dt + 1 day)
    tbl1_jpg            : JPEG bytes — Table 1 (original / before action)
    tbl2_jpg            : JPEG bytes — Table 2 (corrected / after action)
    excel_screenshot_jpg: Optional JPEG — inserted full-width on page 4
    weight_label        : Label for page 5 heading (YES decision only);
                          empty string → page 5 is omitted (NO decision)

    Returns
    -------
    (docx_bytes: bytes, filename: str)
        docx_bytes — the complete .docx file as raw bytes
        filename   — suggested save filename
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Reset image counter for each fresh build
    _IMG_CTR[0] = 0

    doc = Document()
    # Remove the default empty paragraph python-docx adds
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # ── A4 page setup ─────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width       = Cm(21.0)
    sec.page_height      = Cm(29.7)
    sec.left_margin      = Cm(2.54)
    sec.right_margin     = Cm(2.54)
    sec.top_margin       = Cm(2.5)
    sec.bottom_margin    = Cm(2.0)
    sec.header_distance  = Cm(1.27)
    sec.footer_distance  = Cm(1.27)

    # Load logo bytes (silent fail if files are missing)
    toa_b = LOGO_TOA_PATH.read_bytes()     if LOGO_TOA_PATH.exists()     else b""
    sam_b = LOGO_SAMSUNG_PATH.read_bytes() if LOGO_SAMSUNG_PATH.exists() else b""

    _build_headers_footers(doc, rr_line, toa_b, sam_b)

    # ── Convenience shortcuts ─────────────────────────────────────────────────
    def _sp(pts=4):
        """Spacer paragraph."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(pts)

    def _hd(text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
        """Simple one-line heading paragraph."""
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(size)
        r.bold = bold

    # ══════════════════════════════════════════════════════════════════════════
    # PAGE 1: RR number, date, narrative, two weighbridge tables, net weight,
    #         signature block
    # ══════════════════════════════════════════════════════════════════════════

    # RR reference line (underlined)
    p_rr = doc.add_paragraph()
    p_rr.paragraph_format.space_before = Pt(0)
    p_rr.paragraph_format.space_after  = Pt(6)
    rn = p_rr.add_run(rr_line)
    rn.font.name = "Arial"
    rn.font.size = Pt(10)
    rn.underline = True

    _sp(4)
    _add_date_para(doc, "Date: ", rpt_dt)
    _sp(4)

    # Narrative: dynamically selected by reason (A / B / C)
    _add_narrative(
        doc, before_dict, token, arr_dt,
        reason=reason,
    )

    # Small table 1: original weighbridge data (OUT/NET cells yellow)
    _hd("Weighbridge Management System (Net weight)", size=10)
    _sp(6)
    _embed_image(doc, tbl1_jpg)
    _sp(6)

    # Small table 2: corrected data
    _hd("Corrected Net Weight", size=10, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    _sp(6)
    _embed_image(doc, tbl2_jpg)
    _sp(4)

    # "Correct net weight: X" — NO → 0, YES → actual net weight
    # Format: integer "0" when net==0, else decimal string
    accepted_val_p1 = str(action_data.get("ACCEPTED", "")).upper().strip()
    if accepted_val_p1 == "NO":
        nw = "0"   # Action Required = No → Net Weight is always 0
    else:
        nw_raw = action_data.get("NET WEIGHT", "") or ""
        try:
            nw_f = float(nw_raw)
            nw = "0" if nw_f == 0 else (str(int(nw_f)) if nw_f == int(nw_f) else str(round(nw_f, 3)))
        except (ValueError, TypeError):
            nw = nw_raw or "0"
    p_nw = doc.add_paragraph()
    p_nw.paragraph_format.space_before = Pt(0)
    p_nw.paragraph_format.space_after  = Pt(196)   # push sigs to bottom
    r1 = p_nw.add_run("Correct net weight")
    r1.font.name = "Arial"; r1.font.size = Pt(10); r1.bold = True
    r2 = p_nw.add_run(f": {nw}")
    r2.font.name = "Arial"; r2.font.size = Pt(10); r2.bold = True

    _add_signature_block(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # PAGES 2–N: Attachment / reference pages
    #
    # Action Required = YES → 6 pages total (includes Unladen Weight on page 5)
    # Action Required = NO  → 5 pages total (Unladen Weight page omitted)
    #
    # Page layout:
    #   Page 2: Photos of truck no. {VEHICLE}
    #   Page 3: Breakdown truck no. {VEHICLE}
    #   Page 4: Summary Report from Weighbridge Management System dated {DATE}
    #            (optional full-width screenshot embedded here)
    #   Page 5: Unladen Weight heading (YES only)
    #   Page 5/6: Refer Manual Sheet
    # ══════════════════════════════════════════════════════════════════════════

    truck = before_dict.get("VEHICLE NO", "[VEH]")
    d_sfx = f"{arr_dt.day}{_sfx(arr_dt.day)}{arr_dt.strftime(' %b %Y')}"
    out_w_raw = str(action_data.get("OUT WEIGHT", "") or "").strip()

    # ── Page 5 title is REASON-driven (not a UI dropdown anymore) ────────────
    # Option A – Accepted Towing Vehicle  → "Refer Unladen Weight"
    # Option B – Rejected Towing Vehicle  → "Refer Unladen Weight"
    # Option C – Late Time / Breakdown    → "Out Weight from Weighbridge Indicator: X T"
    reason_norm = str(reason).strip()
    if reason_norm == REASON_C:
        # Option C: show actual Out Weight value
        if out_w_raw:
            page5_title = f"Out Weight from Weighbridge Indicator: {out_w_raw} T"
        else:
            page5_title = "Out Weight from Weighbridge Indicator"
    else:
        # Option A or B: always Refer Unladen Weight
        page5_title = "Refer Unladen Weight"

    # Page 5 (weight reference page) is included ONLY when Action Required = YES.
    # Reason (A/B/C) affects CONTENT only, never the page count.
    accepted_val = str(action_data.get("ACCEPTED", "")).upper().strip()
    include_page5 = (accepted_val == "YES")   # strictly YES/NO controlled

    if include_page5:
        page_titles = [
            f"Photos of truck no. {truck}",
            f"Breakdown truck no. {truck}",
            f"Summary Report from Weighbridge Management System dated {d_sfx}",
            page5_title,
            "Refer Manual Sheet",
        ]
    else:
        # Rejected (NO) with Option A or B: 4 attachment pages
        page_titles = [
            f"Photos of truck no. {truck}",
            f"Breakdown truck no. {truck}",
            f"Summary Report from Weighbridge Management System dated {d_sfx}",
            "Refer Manual Sheet",
        ]

    for i, title in enumerate(page_titles):
        _add_section_break(doc, rr_line)

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(8)
        r = p.add_run(title)
        r.font.name = "Arial"
        r.font.size = Pt(11)

        # Page 4 (index 2): insert optional full-width screenshot
        if i == 2 and excel_screenshot_jpg:
            _embed_image(doc, excel_screenshot_jpg, full_width=True)

        # Page 5 (YES, index 3): heading line is the complete content
        # Page 5/6 ("Refer Manual Sheet"): intentionally blank body

    # ── Save, patch, return bytes ─────────────────────────────────────────────
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        doc.save(tmp_path)
        _patch_docx(tmp_path)
        docx_bytes = Path(tmp_path).read_bytes()
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass

    b_code   = _MONTH_B.get(arr_dt.month, 44)
    # ── Filename: RR-B-<BCode>-<Year>-<Serial>-<EToken>-<DDMMYYYY>.docx ───
    safe_tok = re.sub(r"[^\w-]", "", str(token)).strip("-")

    # Use the filter date (selected by user in UI) for month/year/date components.
    # Fallback to arr_dt if filter_date_str is empty or cannot be parsed.
    fn_dt = arr_dt  # fallback
    if filter_date_str:
        try:
            from datetime import date as _date
            _d = _date.fromisoformat(filter_date_str.strip())  # YYYY-MM-DD
            fn_dt = datetime(_d.year, _d.month, _d.day, arr_dt.hour, arr_dt.minute)
        except (ValueError, TypeError):
            fn_dt = arr_dt  # keep fallback

    b_code_fn = _MONTH_B.get(fn_dt.month, 44)
    filename = (
        f"RR-B-{b_code_fn}-{fn_dt.year}"
        f"-{rr_serial.strip()}"
        f"-{safe_tok}"
        f"-{fn_dt.strftime('%d%m%Y')}.docx"
    )
    return docx_bytes, filename


# ══════════════════════════════════════════════════════════════════════════════
#  STREAMLIT TRIGGER  —  optional UI entry point (import streamlit at runtime)
# ══════════════════════════════════════════════════════════════════════════════

def render_rr_trigger(
    token:       str,
    source:      str,
    df:          pd.DataFrame,
    row=None,
    action_data: "dict | None" = None,
    excel_bytes: "bytes | None" = None,
) -> None:
    """
    Streamlit UI entry point for the Rectification Report.

    Behaviour:
      1. User enters RR Serial Number in a text input.
      2. Report is generated AUTOMATICALLY (no extra button click needed).
      3. The finished bytes are stored in st.session_state["_rr_latest"]
         so the calling page can expose a download button.

    Parameters
    ----------
    token       : E-Token string (from the Action Required expander)
    source      : "online" or "wb"
    df          : main weighbridge DataFrame
    row         : optional pd.Series (pre-selected row); if None, token
                  is used to look up the row in df
    action_data : dict of operator-entered corrections
                  (keys: ACCEPTED, OUT WEIGHT, NET WEIGHT, UNLADEN WEIGHT)
    excel_bytes : optional raw Excel bytes (stored in session_state for
                  potential later use)
    """
    import streamlit as st

    if not str(token).strip():
        return

    if excel_bytes is not None:
        st.session_state["excel_buf"] = excel_bytes

    action_data = action_data or {}

    # ── Fetch row data ────────────────────────────────────────────────────────
    try:
        if row is not None and not (isinstance(row, pd.Series) and row.empty):
            before_dict = fetch_row_from_series(row, source=source)
            if not before_dict.get("TOKEN"):
                before_dict["TOKEN"] = str(token).strip()
        else:
            before_dict = fetch_row_by_token(df, token, source=source)
    except Exception as e:
        st.error(f"❌  Cannot fetch data for E-Token [{token}]: {e}")
        return

    arr_dt = _parse_dt(before_dict.get("DATETIME ARRIVAL") or "")
    rpt_dt = arr_dt + timedelta(days=1)
    b_code = _MONTH_B.get(arr_dt.month, 44)

    after_dict      = apply_user_updates(before_dict, action_data)
    accepted_status = str(action_data.get("ACCEPTED", "")).upper()

    # ── RR Serial Number input ────────────────────────────────────────────────
    _sk    = f"rr_serial_{token}"
    serial = st.text_input(
        "RR Serial Number",
        value=st.session_state.get(_sk, ""),
        placeholder="e.g.  0290",
        key=f"rr_serial_inp_{token}",
    )
    if serial.strip():
        st.session_state[_sk] = serial.strip()

    if not serial.strip():
        st.info("ℹ️  Enter the RR Serial Number above to generate the report.")
        return

    rr_line  = f"Rectification Report No. RR/B-{b_code}/{arr_dt.year}/{serial.strip()}"
    safe_tok = re.sub(r"[^\w]", "", str(token))
    filename = (
        f"RR-B-{b_code}-{arr_dt.year}-{serial.strip()}"
        f"-{safe_tok}-{arr_dt.strftime('%d%m%Y')}.docx"
    )
    st.markdown(f"**RR Number:** `{rr_line}`")

    # ── Weight label (YES only) ───────────────────────────────────────────────
    out_w_str    = str(action_data.get("OUT WEIGHT", "") or "").strip()
    _wl_key      = f"rr_weight_label_{token}"
    weight_label = ""

    if accepted_status == "YES" and out_w_str:
        action_data["UNLADEN WEIGHT"] = out_w_str

        wl_display = [
            "Refer Unladen Weight",
            "Out Weight Source (Weighbridge Indicator)",
        ]
        wl_values = [
            "Refer Unladen Weight",
            f"Out Weight from Weighbridge Indicator: {out_w_str} T",
        ]
        _wl_idx_key = f"rr_weight_label_idx_{token}"
        prev_idx    = st.session_state.get(_wl_idx_key, 0)
        sel_idx = st.selectbox(
            "Weight Source for Page 5",
            options=list(range(len(wl_display))),
            format_func=lambda i: wl_display[i],
            index=prev_idx,
            key=f"wl_sel_{token}",
            help="Select how the Unladen Weight sheet heading should appear",
        )
        st.session_state[_wl_idx_key] = sel_idx
        weight_label = wl_values[sel_idx]
        st.session_state[_wl_key] = weight_label

    # ── Auto-generate whenever inputs change ──────────────────────────────────
    _gk       = f"rr_doc_{token}_{serial.strip()}_{weight_label[:20]}"
    _prev_key = f"rr_prev_key_{token}"
    _cur_key  = f"{serial.strip()}|{weight_label}"
    key_changed  = st.session_state.get(_prev_key, "") != _cur_key
    need_generate = _gk not in st.session_state or key_changed

    if need_generate:
        st.session_state[_prev_key] = _cur_key
        with st.spinner("⏳  Building Rectification Report…"):
            try:
                tbl1_jpg, _ = fetch_and_generate(
                    df=df, token=token, source=source,
                    is_table2=False, override_values=None, dpi=300,
                    force_outnet_yellow=True,
                )
                tbl2_jpg, _ = fetch_and_generate(
                    df=df, token=token, source=source,
                    is_table2=True, override_values=action_data, dpi=300,
                )
                docx_bytes, _ = build_rr_docx(
                    token        = token,
                    rr_serial    = serial.strip(),
                    rr_line      = rr_line,
                    before_dict  = before_dict,
                    after_dict   = after_dict,
                    action_data  = action_data,
                    arr_dt       = arr_dt,
                    rpt_dt       = rpt_dt,
                    tbl1_jpg     = tbl1_jpg,
                    tbl2_jpg     = tbl2_jpg,
                    excel_screenshot_jpg = None,   # page 4 screenshot removed
                    weight_label         = weight_label,
                )
                st.session_state[_gk] = {"bytes": docx_bytes, "filename": filename}
                st.session_state["_rr_latest"] = st.session_state[_gk]

            except Exception as ex:
                import traceback
                st.error(f"❌  Failed to generate RR: {ex}")
                st.code(traceback.format_exc())
                return

    # Finished — bytes are in st.session_state["_rr_latest"]
    # The parent app.py exposes the download button.


# ══════════════════════════════════════════════════════════════════════════════
#  STANDALONE TEST  (python rectification_report.py)
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("Running standalone Rectification Report test…")

    # Minimal synthetic row — replace with real data when integrating
    sample_row = {
        "TOKEN":            "APSGHKKL102202604290018",
        "SITE CODE":        "APSGHKKL102",
        "DATETIME ARRIVAL": "29/04/2026 17:14",
        "VEHICLE NO":       "XE3059B",
        "MATERIAL":         "SOFT CLAY",
        "ACCEPTED":         "YES",
        "IN WEIGHT":        "29.91",
        "OUT WEIGHT":       "",
        "NET WEIGHT":       "",
    }
    action = {
        "ACCEPTED":        "YES",
        "OUT WEIGHT":      "10.50",
        "NET WEIGHT":      "19.41",
        "UNLADEN WEIGHT":  "10.50",
    }

    arr_dt = _parse_dt(sample_row["DATETIME ARRIVAL"])
    rpt_dt = arr_dt + timedelta(days=1)
    b_code = _MONTH_B.get(arr_dt.month, 44)
    serial = "0001"
    rr_line = f"Rectification Report No. RR/B-{b_code}/{arr_dt.year}/{serial}"

    before_dict = sample_row.copy()
    after_dict  = apply_user_updates(before_dict, action)

    print("  Generating Table 1…")
    tbl1_jpg = generate_table_image(before_dict, is_table2=False, dpi=300,
                                    force_outnet_yellow=True)
    print(f"  ✅ Table 1: {len(tbl1_jpg):,} bytes")

    print("  Generating Table 2…")
    tbl2_jpg = generate_table_image(after_dict, is_table2=True, dpi=300)
    print(f"  ✅ Table 2: {len(tbl2_jpg):,} bytes")

    print("  Building .docx…")
    docx_bytes, filename = build_rr_docx(
        token        = sample_row["TOKEN"],
        rr_serial    = serial,
        rr_line      = rr_line,
        before_dict  = before_dict,
        after_dict   = after_dict,
        action_data  = action,
        arr_dt       = arr_dt,
        rpt_dt       = rpt_dt,
        tbl1_jpg     = tbl1_jpg,
        tbl2_jpg     = tbl2_jpg,
        weight_label = "Refer Unladen Weight",
    )
    out = Path(filename)
    out.write_bytes(docx_bytes)
    print(f"  ✅ Saved → {out}  ({len(docx_bytes):,} bytes)")
    print("Done.")
